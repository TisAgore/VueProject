import fitz
from docx import Document
from openpyxl import Workbook
from PIL import Image

from auditmate.tools.parsers import parse_attachment


def test_parse_text_and_json(tmp_path):
    text_path = tmp_path / "notes.md"
    text_path.write_text("Runway: 12 months", encoding="utf-8")
    json_path = tmp_path / "data.json"
    json_path.write_text('{"ARR": 1200000}', encoding="utf-8")

    text_parsed = parse_attachment(text_path)
    json_parsed = parse_attachment(json_path)

    assert "Runway" in text_parsed.text
    assert json_parsed.data["json"]["ARR"] == 1200000


def test_parse_csv_extracts_financial_metrics(tmp_path):
    path = tmp_path / "metrics.csv"
    path.write_text("metric,value\nARR,1200000\nburn rate,50000\n", encoding="utf-8")

    parsed = parse_attachment(path)

    assert parsed.data["financial_metrics"]["arr"] == "1200000"
    assert parsed.data["financial_metrics"]["burn_rate"] == "50000"


def test_parse_xlsx_extracts_financial_metrics(tmp_path):
    path = tmp_path / "metrics.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["metric", "value"])
    sheet.append(["ARR", 1200000])
    sheet.append(["runway", 18])
    workbook.save(path)
    workbook.close()

    parsed = parse_attachment(path)

    assert parsed.data["financial_metrics"]["arr"] == 1200000
    assert parsed.data["financial_metrics"]["runway_months"] == 18


def test_parse_pdf_extracts_text(tmp_path):
    path = tmp_path / "pitch.pdf"
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Revenue 1000")
    document.save(path)
    document.close()

    parsed = parse_attachment(path)

    assert "Revenue 1000" in parsed.text
    assert parsed.metadata["pages"] == 1
    assert parsed.data["image_page_previews"] == []


def test_parse_image_based_pdf_renders_page_images(tmp_path):
    path = tmp_path / "scan.pdf"
    document = fitz.open()
    page = document.new_page()
    page.draw_rect((72, 72, 240, 160), color=(0, 0, 0), fill=(1, 1, 1))
    document.save(path)
    document.close()

    parsed = parse_attachment(path)

    assert parsed.data["image_based"] is True
    assert parsed.data["page_images"]
    assert parsed.data["page_images"][0]["mime_type"] == "image/png"


def test_parse_pdf_collects_pages_with_embedded_images(tmp_path):
    image_path = tmp_path / "embedded.png"
    Image.new("RGB", (20, 20), "black").save(image_path)

    path = tmp_path / "deck.pdf"
    document = fitz.open()
    page1 = document.new_page()
    page1.insert_text((72, 72), "Revenue 1000")
    page2 = document.new_page()
    page2.insert_image(fitz.Rect(72, 72, 172, 172), filename=str(image_path))
    document.save(path)
    document.close()

    parsed = parse_attachment(path)

    assert parsed.data["image_page_numbers"] == [2]
    assert len(parsed.data["image_page_previews"]) == 1
    assert parsed.data["image_page_previews"][0]["page"] == 2


def test_parse_docx_extracts_text(tmp_path):
    path = tmp_path / "memo.docx"
    document = Document()
    document.add_heading("Team", level=1)
    document.add_paragraph("CEO: Test Founder")
    document.save(path)

    parsed = parse_attachment(path)

    assert "CEO: Test Founder" in parsed.text
    assert "Team" in parsed.data["headings"]


def test_parse_image_extracts_metadata(tmp_path):
    path = tmp_path / "chart.png"
    Image.new("RGB", (32, 16), "white").save(path)

    parsed = parse_attachment(path)

    assert parsed.metadata["width"] == 32
    assert parsed.metadata["height"] == 16
    assert parsed.data["base64"]
