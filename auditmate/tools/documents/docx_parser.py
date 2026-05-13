from pathlib import Path

from auditmate.models.attachments import Attachment, ParsedAttachment, create_attachment


class DOCXParserTool:
    name = "parse_docx"
    description = "Извлекает текст, заголовки и таблицы из DOCX-файлов."

    def execute(self, input_data: Attachment | str | Path) -> ParsedAttachment:
        """Парсит DOCX-файл через python-docx."""

        attachment = _coerce_attachment(input_data)
        try:
            from docx import Document
        except ImportError as e:
            raise RuntimeError(
                "Для чтения DOCX-файлов установите зависимость: python-docx."
            ) from e

        document = Document(str(attachment.path))
        paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
        headings = [
            p.text.strip()
            for p in document.paragraphs
            if p.text.strip() and p.style and p.style.name.lower().startswith("heading")
        ]
        tables = []
        for table in document.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if rows:
                tables.append(rows)

        table_text = _tables_to_text(tables)
        text_parts = paragraphs + ([table_text] if table_text else [])
        text = "\n".join(text_parts)

        return ParsedAttachment(
            attachment=attachment,
            text=text,
            data={
                "headings": headings,
                "paragraphs": paragraphs,
                "tables": tables,
            },
            metadata={
                "source": str(attachment.path),
                "kind": attachment.kind,
                "paragraphs": len(paragraphs),
                "tables": len(tables),
                "characters": len(text),
            },
        )


def _tables_to_text(tables: list[list[list[str]]]) -> str:
    """Преобразует таблицы DOCX в компактный текстовый блок."""

    blocks = []
    for table_index, table in enumerate(tables, start=1):
        rows = [" | ".join(cell for cell in row if cell) for row in table]
        blocks.append(f"Таблица {table_index}:\n" + "\n".join(row for row in rows if row))

    return "\n\n".join(blocks)


def _coerce_attachment(input_data: Attachment | str | Path) -> Attachment:
    """Нормализует вход парсера в объект Attachment."""

    if isinstance(input_data, Attachment):
        return input_data

    return create_attachment(input_data)

